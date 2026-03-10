import nltk
from nltk.util import ngrams
from collections import Counter
import pruebaTokenExcel
from docx import Document

tokens = pruebaTokenExcel.allTokens

# Generar Ngramas
n = 2  # tamaño
bigrams = list(ngrams(tokens, n))

# Contar la frecuencia de los N-gramas
frecuencia_bigrams = Counter(bigrams)

# Ordenar por frecuencia
bigrams_mas_comunes = frecuencia_bigrams.most_common(20)  # Top X más comunes

# Obtener todos los Ngramas 
todos_bigrams = frecuencia_bigrams.most_common()  # Ordenados por frecuencia descendente

# Ruta para guardar el archivo Word
output_path = 'C:\\tfg_sergio\\word_test_ngram.docx'

# Crear un documento de Word
doc = Document()

# Título del documento
doc.add_heading(f'Todos los {n}-gramas ordenados por frecuencia', level=1)

# Escribir todos los Ngramas y sus frecuencias en el documento
for ngram, freq in todos_bigrams:

    paragraph = doc.add_paragraph()
    paragraph.add_run("{n}-gram: ")
    
    # Añadir el bigrama con la fuente Iberian
    ngram_text = ' '.join(ngram)
    run = paragraph.add_run(ngram_text)
    font = run.font
    font.name = 'iberian'  # Aplicar la fuente Iberian a los bigramas
    
    # Añadir la frecuencia con una fuente estándar
    paragraph.add_run(f" - Frecuencia: {freq}")

# Guardar el archivo de Word
doc.save(output_path)

print(f"Todos los bigramas y frecuencias exportados a {output_path}")

