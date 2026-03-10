import pandas as pd
from collections import Counter
from docx import Document
import re
import os

# Ruta del Excel
path = os.path.dirname(os.path.abspath(__file__))
file_path = path + '\ibers_febrer_24_maria.xls' #'\ibers_febrer_24.xls'
#outputPath = 'C:\\Users\\sergi\\Desktop\\ArchivosTFG\\FrecuenciasTokens.docx'

#lista para guardar los tokens
allTokens = []

# Lee el Excel
df = pd.read_excel(file_path)

# Muestra las primeras filas del DataFrame para verificar la lectura
print(df.head())

# Procesa los datos 
for index, row in df.iterrows():
    referencia = row[0]
    contenido_inscripcion = row[1]
    tipo_inscripcion = row[2]
    comentario = row[3]
    fuenteArial = row[4]

    
    #print(f" insc {index+1}: {fuenteArial}")

    #Tokenización
    if isinstance(fuenteArial, str):  #Verifica si el contenido es una cadena
        fuenteArial = fuenteArial.strip()  #elimina espacios al inicio y al final
        if fuenteArial:  #Verifica si no está vacío después de limpiar
            #tokens = fuenteArial.split()  #Divide la cadena por espacios
            tokens = re.split(r'[ \u00f9\u00f6\u00f4\u2192]', fuenteArial)
            tokens = [token for token in tokens if token]  # Elimina tokens vacíos
            print(f" insc {index+1}, Tokens: {tokens}")
            allTokens.extend(tokens)  #Agrega los tokens a la lista general
        else:
            print(f" insc {index+1}, Tokens: Contenido vacío")
    else:
        print(f" insc {index+1}, Tokens: No hay contenido de texto para tokenizar")

        
#Calcula las frecuencias de los tokens
tokenFrequencies = Counter(allTokens)

filteredTokens = sorted(
    {token: freq for token, freq in tokenFrequencies.items() if freq >= 2}.items(),
    key=lambda x: x[1],  # Ordena por la frecuencia (valor)
    reverse=True  # Orden descendente
)

#Imprime las frecuencias de los tokens
print("\nFrecuencia de los tokens:")
for token, freq in filteredTokens:
    print(f"Token: '{token}' - Frecuencia: {freq}")


outputPath = path + "\word_test.docx"

# Crea un documento de Word
doc = Document()

# Título del documento
doc.add_heading('Tokens con frecuencia mayor o igual a 2, ordenados de mayor a menor:', level=1)

# Escribe los tokens y sus frecuencias en el documento
'''for token, freq in filteredTokens:
    doc.add_paragraph(f"Token: '{token}' - Frecuencia: {freq}")
'''
for token, freq in filteredTokens:
    paragraph = doc.add_paragraph() # EN vez de escribir todo de golpe, dividimos en tres partes
    paragraph.add_run(f"Token: ")
    run2 = paragraph.add_run(f'{token}')
    font2 = run2.font # A la segunda parte le estipulamos "iberian" como font
    font2.name = 'iberian'
    paragraph.add_run(f" - Frecuencia: {freq}")


# Guarda el documento en un archivo .docx
doc.save(outputPath)

print(f"Frecuencias de tokens guardadas en: {outputPath}")
