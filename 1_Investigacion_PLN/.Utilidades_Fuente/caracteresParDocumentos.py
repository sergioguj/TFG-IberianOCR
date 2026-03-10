import string
from docx import Document
from docx.shared import Pt

# Configuración de la fuente y el tamaño
font_name = "Iberian"
font_size = 14

# Definir los caracteres a incluir
characters = string.ascii_letters + string.digits + string.punctuation + "¡¿"

# Iterar sobre cada carácter y generar los archivos correspondientes
for index, char in enumerate(characters, start=1):
    # Crear el archivo de texto con codificación UTF-8
    txt_filename = f"caracter_{index}.txt"
    with open(txt_filename, "w", encoding="utf-8") as txt_file:
        txt_file.write(char)
    
    # Crear el documento Word para el carácter
    doc = Document()
    # Se añade un párrafo con el carácter
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(char)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    
    docx_filename = f"caracter_{index}.docx"
    doc.save(docx_filename)
    
    print(f"Archivos generados para el carácter '{char}': {txt_filename} y {docx_filename}")
