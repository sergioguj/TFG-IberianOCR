from docx import Document
from docx.shared import Pt

# Crear un nuevo documento
doc = Document()

# Configuración de la fuente y el tamaño
font_name = "Iberian"
font_size = 14

# Caracteres que queremos incluir
import string
characters = string.ascii_letters + string.digits + string.punctuation + "¡¿"

# Añadir título al documento
doc.add_heading("Caracteres de la fuente Iberian", level=1)

# Añadir los caracteres al documento
paragraph = doc.add_paragraph()
run = paragraph.add_run("Caracteres disponibles:\n")
run.bold = True

for char in characters:
    run = paragraph.add_run(char + " ")  # Añade cada carácter con un espacio
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)

# Guardar el documento
output_path = "caracteres_iberian.docx"
doc.save(output_path)

print(f"Documento generado: {output_path}")
